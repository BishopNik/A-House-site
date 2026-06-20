from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING


OUT = "/Users/bishop/Documents/A-House-site/A-House_1_Core_Manufacturer_Design_Brief_EN.docx"
CONCEPT_IMAGE = "/Users/bishop/Documents/A-House-site/A-House_1_Core_original_single_storey_concept.png"

NAVY = "17324D"
BLUE = "2E74B5"
MID = "526578"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GRID = "CBD5DF"
WHITE = "FFFFFF"
BLACK = "111111"
GOLD = "A16A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=BLACK, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, style=None, bold_lead=None, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullet(doc, text, level=0, bold_lead=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return add_text(doc, text, style=style, bold_lead=bold_lead, after=4)


def add_number(doc, text):
    return add_text(doc, text, style="List Number", after=4)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12, 4: 11}[level], bold=True,
            color=BLUE if level < 4 else NAVY)
    return p


def add_callout(doc, label, text, fill=PALE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label + "  ")
    set_run(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run(r, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_run(r, size=9, color=MID)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for i, spec in {1: (16, 18, 10), 2: (13, 14, 7), 3: (12, 10, 5), 4: (11, 8, 4)}.items():
        name = f"Heading {i}"
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st = styles[name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.font.size = Pt(spec[0])
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if i < 4 else NAVY)
        st.paragraph_format.space_before = Pt(spec[1])
        st.paragraph_format.space_after = Pt(spec[2])
        st.paragraph_format.keep_with_next = True
    for name, left, first in (("List Bullet", 0.375, -0.188), ("List Bullet 2", 0.68, -0.18), ("List Number", 0.375, -0.188)):
        st = styles[name]
        st.font.name = "Aptos"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(left)
        st.paragraph_format.first_line_indent = Inches(first)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_section_intro(doc, title, purpose):
    add_heading(doc, title, 2)
    add_text(doc, purpose, after=7, keep=True)


doc = Document()
configure_styles(doc)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

# Running header/footer
hp = section.header.paragraphs[0]
hp.text = ""
hp.paragraph_format.space_after = Pt(0)
rh = hp.add_run("A-HOUSE 1 CORE  |  MANUFACTURER DESIGN BRIEF")
set_run(rh, size=8.5, bold=True, color=MID)
add_page_number(section.footer.paragraphs[0])

# Customer-pack opening block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("SUPPLIER DESIGN & QUOTATION PACKAGE")
set_run(r, size=10, bold=True, color=GOLD)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("A-House 1 Core")
set_run(r, size=30, bold=True, color=NAVY, font="Aptos Display")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Manufacturer Design Brief for Prefabricated House Suppliers in China")
set_run(r, size=14, color=MID)

image_p = doc.add_paragraph()
image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
image_p.paragraph_format.space_after = Pt(4)
image_run = image_p.add_run()
inline_shape = image_run.add_picture(CONCEPT_IMAGE, width=Inches(6.5))
inline_shape._inline.docPr.set("descr", "Original single-storey A-House concept image from the source design brief.")
caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(10)
r = caption.add_run("Original single-storey concept from the source brief — visual direction only; not a final design or construction detail.")
set_run(r, size=9, italic=True, color=MID)

meta = doc.add_table(rows=2, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(meta, [4680, 4680], indent=120)
meta_data = [
    ("DOCUMENT PURPOSE", "Concept design, engineering proposal and comparative quotation"),
    ("STATUS", "Design brief — not for construction"),
    ("PRIMARY CONCEPT", "Single-storey rectangular house"),
    ("UNITS / LANGUAGE", "Metric (mm, m, °C, W, kW) / English"),
]
for idx, (label, value) in enumerate(meta_data):
    cell = meta.cell(idx // 2, idx % 2)
    set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label + "\n")
    set_run(r, size=8.5, bold=True, color=BLUE)
    r = p.add_run(value)
    set_run(r, size=10.5)

doc.add_page_break()
doc.add_paragraph().paragraph_format.space_after = Pt(3)
add_callout(
    doc,
    "PRIMARY DESIGN DIRECTION",
    "The first concept shall be developed as a single-storey house. A two-storey version is acceptable only as a separately priced option if the total price increase is minor and clearly demonstrated by the supplier.",
    fill="FFF8E8", accent=GOLD,
)
add_callout(
    doc,
    "FOUNDATION DESIGN BASIS",
    "The pile-and-grade-beam foundation is a preliminary baseline for quotation. The final foundation may change after the site survey, geotechnical investigation and written recommendations of the responsible structural designer. No foundation solution may be released for construction before local verification.",
)

add_heading(doc, "1. Purpose and Supplier Response", 1)
add_text(doc, "This brief defines the owner’s functional, architectural and technical intent for a highly energy-efficient, factory-manufactured modular house. The supplier shall convert this intent into a coordinated concept, an engineering proposal, a production strategy and a transparent quotation.")
add_text(doc, "The supplier shall identify every deviation, assumption, exclusion and item requiring local engineering approval. Silence shall not be treated as acceptance of a requirement.")
add_heading(doc, "Required supplier submission", 2)
for item in [
    "Single-storey concept package: plans, elevations, sections, key dimensions, gross/internal areas and preliminary module breakdown.",
    "Optional two-storey concept and separate price comparison, including the absolute and percentage price difference versus the single-storey baseline.",
    "Structural concept with load paths, module joints, lifting points, transport restraints and site-assembly sequence.",
    "Foundation interface loads and reactions; final foundation design to be coordinated with local survey and soil data.",
    "Envelope build-ups, calculated thermal performance, air-tightness strategy, condensation-risk review and thermal-bridge details.",
    "MEP concept diagrams, equipment schedule, access zones and maintenance/replacement strategy.",
    "Bill of materials, product data sheets, certificates, fire/corrosion/durability information and proposed equivalents.",
    "Factory scope, site scope, exclusions, lead time, packing method, container/loading plan, assembly crew requirements and installation manual.",
    "Itemised quotation with currency, Incoterms, validity period, payment milestones, warranties, spare parts and optional items.",
    "Compliance matrix referencing every requirement in this brief and marking Comply / Alternative / Excluded / Information Required.",
]: add_bullet(doc, item)

add_heading(doc, "2. Design Priorities and Non-Negotiable Principles", 1)
priorities = [
    ("Energy efficiency.", "Maintain a continuous thermal and airtight envelope. Unjustified heat-loss paths, thermal bridges, cold thresholds and hidden condensation risks are not acceptable."),
    ("Fast assembly.", "Use repeatable factory-made modules that are transportable and can be assembled rapidly on site. Base the dimensional grid on the 1,220 mm MgO board module where practical."),
    ("Repairability.", "All structural and MEP systems must permit local inspection, servicing, replacement and repair without demolition of large building areas."),
    ("Planning flexibility.", "Avoid internal load-bearing posts in the living volume. Internal partitions shall be modular, demountable and acoustically insulated."),
    ("Simple site work.", "Minimise the number and specialist skill level of site workers. Provide detailed assembly drawings, instructions and quality-control hold points for each stage."),
    ("Supplier accountability.", "All final dimensions, materials, joints and capacities shall be supported by calculations and approved by the responsible designers for the manufacturing and construction jurisdictions."),
]
for lead, body in priorities: add_bullet(doc, lead + " " + body, bold_lead=lead)

add_heading(doc, "3. Key Project Decisions", 1)
key_decisions = [
    "First concept: single-storey rectangular house. Thermal SIP envelope: 15 mm MgO / 220 mm PIR / 15 mm MgO as the target baseline, subject to verified performance and structural design.",
    "Optional concept: two storeys only if the incremental price is minor; quote and explain the cost delta separately.",
    "Preliminary foundation: pile field with pile caps, primary beams and grade beams. Final system may change following site/geodetic survey, geotechnical results and structural designer recommendations.",
    "Service crawlspace: enclosed, warm, insulated and accessible; no uncontrolled outside-air flushing. Floor graded to an emergency drain.",
    "Roof void: warm, accessible service space inside the thermal and airtight envelopes; not a cold ventilated attic.",
    "Façade: modular, drained, ventilated and locally replaceable. Solid façade elements use paint-grade fibre-cement boards. All penetrative supports require thermal breaks.",
    "Roof: warm low-slope mono-pitch appearance, target slope approximately 4°, with its own complete waterproofing independent of the PV superstructure.",
    "PV array: separate maintainable superstructure above the roof, with ventilated under-panel space, safe service walkways, cable management and drainage.",
    "Windows and external doors: high-performance double- or triple-chamber insulating glass units, thermally broken frames, insulated support profiles/thresholds and three-layer installation joints.",
    "Horizontal drainage and other service routes remain visible and accessible in the crawlspace or designated service zones. Concealed non-serviceable connections are prohibited.",
    "Primary heating/cooling: heat pump, preferably ground-source subject to site feasibility. Dry hydronic floor heating and dry hydronic radiant ceiling cooling with dew-point protection.",
    "Ventilation: balanced supply/extract system with heat recovery, controlled dehumidification and dew-point monitoring.",
    "Backup heating: pellet boiler. Domestic hot water: solar thermal collectors and storage cylinder, with heat-pump and pellet-boiler backup.",
    "Electrical concept: PV-led generation with grid connection as backup. Include 230/400 V distribution as required by the destination market, low-voltage systems and EV charging; confirm final voltage/frequency with the owner.",
]
for item in key_decisions: add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "4. Basic Design Solution Register", 1)
add_text(doc, "The register coordinates the complete project scope. Detailed requirements currently included in this brief are marked “Detailed”. Remaining disciplines shall be developed and cross-coordinated during the next design stage.")
register = [
    (1,"House","BDS-01","Foundation","Detailed"),(2,"House","BDS-02","Service crawlspace","Detailed"),(3,"House","BDS-03","Insulated perimeter apron","Detailed"),(4,"House","BDS-04","Façade","Detailed"),(5,"House","BDS-05","Living floor(s)","Detailed"),(6,"House","BDS-06","Warm roof service void","Detailed"),(7,"House","BDS-07","Roof","Detailed"),(8,"House","BDS-08","PV superstructure","Detailed"),(9,"House","BDS-09","External stair tower","Optional / detailed"),(10,"House","BDS-10","External doors and windows","Detailed"),(11,"House","BDS-11","House thermal envelope","Detailed"),
    (12,"Interior","BDS-12","Serviceable heated-floor build-up","Detailed"),(13,"Interior","BDS-13","Serviceable radiant cooling ceiling","Detailed"),(14,"Interior","BDS-14","Internal partitions, doors and windows","Detailed"),(15,"Interior","BDS-15","Internal lining to external walls","Detailed"),
    (16,"Site","BDS-16","Car canopy","Coordinate"),(17,"Site","BDS-17","Plant room","Coordinate"),(18,"Site","BDS-18","Paths","Coordinate"),(19,"Site","BDS-19","Fence","Coordinate"),
    (20,"Heating/cooling","BDS-20","Heat and cooling generation","Coordinate"),(21,"Heating/cooling","BDS-21","Heated-floor hydraulics","Coordinate"),(22,"Heating/cooling","BDS-22","Cooling-ceiling hydraulics","Coordinate"),(23,"Heating/cooling","BDS-23","Ventilation","Coordinate"),(24,"Heating/cooling","BDS-24","Dehumidification","Coordinate"),
    (25,"Electrical","BDS-25","Grid / main switchboard","Coordinate"),(26,"Electrical","BDS-26","Solar PV system","Coordinate"),(27,"Electrical","BDS-27","ELV, 230 V and 400 V distribution","Coordinate"),(28,"Electrical","BDS-28","EV charging","Coordinate"),(29,"Electrical","BDS-29","Earthing and lightning protection","Coordinate"),(30,"Electrical","BDS-30","House lighting","Coordinate"),(31,"Electrical","BDS-31","Site lighting","Coordinate"),
    (32,"Pellet boiler","BDS-32","Snow-melting system","Coordinate"),(33,"Pellet boiler","BDS-33","Hammam steam generation","Coordinate"),(34,"Pellet boiler","BDS-34","Backup heating","Coordinate"),
    (35,"Water","BDS-35","Cold water supply","Coordinate"),(36,"Water","BDS-36","Domestic hot water","Coordinate"),(37,"Water","BDS-37","Hammam water supply","Coordinate"),(38,"Water","BDS-38","Irrigation","Coordinate"),
    (39,"Other services","BDS-39","Stormwater drainage","Detailed"),(40,"Other services","BDS-40","Site drainage","Coordinate"),(41,"Other services","BDS-41","Sanitary drainage","Coordinate"),(42,"Other services","BDS-42","Alarm and CCTV","Coordinate"),(43,"Other services","BDS-43","Internet / Wi-Fi / ELV networks","Coordinate"),
]
t = doc.add_table(rows=1, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(t, [700, 1600, 950, 4210, 1900], indent=120)
headers = ["No.", "Discipline", "Ref.", "Title", "Current status"]
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]; set_cell_shading(c, LIGHT)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,2,4) else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(h); set_run(r, size=9, bold=True, color=NAVY)
set_repeat_table_header(t.rows[0])
for row in register:
    cells = t.add_row().cells
    for i, val in enumerate(row):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,2,4) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
        r = p.add_run(str(val)); set_run(r, size=8.5, color=BLACK)
        if row[0] % 2 == 0: set_cell_shading(cells[i], "F8FAFC")

doc.add_page_break()
add_heading(doc, "5. Detailed Basic Design Solutions", 1)

add_section_intro(doc, "BDS-01 — Foundation", "Preliminary basis: pile-and-grade-beam foundation for a lightweight modular house, coordinated with the 1,220 mm MgO module.")
add_callout(doc, "MANDATORY DESIGN HOLD POINT", "Final pile type, material, diameter/section, pile caps, primary beams, grade beams, connections and upper support frame shall be determined from the actual site survey, geotechnical report, loads and responsible structural engineer’s recommendations. The supplier shall provide interface loads, not assume that the preliminary foundation is final.", fill="FFF8E8", accent=GOLD)
for item in [
    "Locate piles at calculated support points below primary beams, grade beams and concentrated loads.",
    "Provide the floor support deck above the grade beams using structural SIP panels; target build-up 15 mm MgO / 200–250 mm PIR / 15 mm MgO, subject to structural and fire verification.",
    "The upper support frame shall receive the first-floor SIP envelope and the crawlspace enclosure.",
    "Maintain an accessible crawlspace clear height of approximately 1,200–1,500 mm from finished crawlspace floor to the underside of the structural floor deck.",
    "The perimeter apron is non-structural and shall not be rigidly connected to piles, beams, grade beams or plinth frame.",
    "Route horizontal services in the crawlspace. All foundation penetrations shall be designed, sleeved, sealed and accessible.",
    "Install underground sanitary drainage and cold-water entries before piling where sequencing requires. Keep pile caps, connections, support frame and service sleeves accessible for inspection.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-02 — Warm Service Crawlspace", "Provide a closed, warm and externally insulated service zone beneath the house, with approximately 1,200 mm clear height.")
for item in [
    "Maintain at least +5 °C at the coldest control points and 40–60% RH. Suggested controls: heating off at +8 °C; alarm at +3 °C or below.",
    "No uncontrolled ventilation or free flushing by outdoor air. Provide a dedicated extract duct independent of the sanitary vent stack.",
    "Grade the entire crawlspace floor gently to an emergency drain. Seal interfaces at piles and service penetrations.",
    "Ground-up floor build-up: continuous sealed vapour/waterproof membrane; two staggered layers of 50–100 mm XPS; optional 40 mm routed XPS with hydronic pipes and aluminium heat-spreader plates; 15 mm MgO service deck with elastomeric sealed joints.",
    "Support the enclosure on a designed plinth frame fixed to the pile/beam foundation—not on soil, XPS, geotextile or the perimeter apron.",
    "Typical wall panel: LVL-edged SIP, 15 mm MgO / at least 200 mm PIR / 15 mm MgO. Add 50 mm external mineral wool over framing to reduce thermal bridges, weather membrane, 20–30 mm drained cavity and fibre-cement protection.",
    "Use 50 mm XPS with rigid protection in the lower 1 m splash zone. Include drip edge, protected edges, insect/rodent mesh and sealed perimeter-apron junction.",
    "Provide sealed external service hatches on all sides as required by the maintenance layout; original target is four per side, to be rationalised during design without reducing access.",
    "Floor-deck joints form part of the airtight layer. No inaccessible services may be concealed within the floor deck.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-03 — Insulated Perimeter Apron", "Provide a continuous insulated apron around the house to protect the base, control surface water and support snow-melting where required.")
for item in [
    "Approximate width: 1,200 mm from the crawlspace enclosure; slope: 3–5% away from the building to channels, drainage or an approved discharge point.",
    "The apron is not a foundation element and shall remain structurally separated from piles, pile caps, beams, grade beams and the plinth frame.",
    "Do not block service sleeves, cleanouts, hatches or exterior maintenance points. Provide a sealed movement joint at the building.",
    "Typical ground-up build-up: compacted graded soil; geotextile ≥200 g/m² with ≥200 mm laps; compacted crushed-stone base; levelling sand; two staggered 50 mm XPS layers; protective separation layer; reinforced concrete slab with glycol snow-melting loop; compatible exterior waterproof/adhesive layer; natural stone finish with frost-resistant flexible joints.",
    "Protect all exposed XPS and lower-wall details against water, UV, impact and rodents.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-04 — Façade", "Provide a ventilated, drained, modular and repairable façade aligned to the 1,220 mm board/module grid where practical.")
for item in [
    "Architectural intent: simple rectangular volume with large light and dark painted fibre-cement panels, restrained colour inserts and optional fins.",
    "Typical sequence: structural/enclosure layer; continuous external insulation over SIP framing; weather-resistive membrane; 20–30 mm drained cavity; subframe; removable external cladding.",
    "Coordinate visible joints with modules and openings; avoid arbitrary visible board cuts.",
    "Verify corrosion compatibility, differential movement, wind, fire and exterior durability. Every bracket, anchor or fastener crossing insulation requires a calculated thermal break.",
    "Horizontal bands may protect the plinth, mark floor lines or conceal the roof/PV edge, but shall not obstruct drainage, ventilation or maintenance.",
    "Protect the crawlspace splash zone and maintain sealed, visually integrated service hatches. Neither façade nor crawlspace enclosure may bear on the perimeter apron.",
    "Fibre-cement fins are permitted only after checking thickness, wind loads, fixings, edge protection, weight, fire, shading, greenhouse ventilation and replaceability.",
    "The stair-tower envelope is structurally independent; its junction shall not transfer vertical loads or admit water.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-05 — Living Floor(s)", "The first concept shall contain one living storey. Design the structural grid and interfaces so that a separately priced two-storey alternative can be evaluated without compromising the baseline concept.")
for item in [
    "Single-storey baseline: internal SIP envelope forms the principal thermal and airtight enclosure; no internal load-bearing posts in living areas.",
    "Where a glazed outer buffer/greenhouse enclosure is proposed, use high-performance glazing in transparent zones and insulated SIP or equivalent construction in opaque zones.",
    "For an optional second storey, use engineered floor trusses or another calculated system without internal posts. Provide an independent external stair tower if required.",
    "Route services through crawlspace, designated floor/ceiling service zones and warm roof void, with access to all joints and inspection points.",
    "Use BDS-14 demountable internal partitions for room planning.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-06 — Warm Roof Service Void", "Provide a warm, maintainable service volume between the living ceiling and roof. For the single-storey concept this is above the first floor; for the optional two-storey concept it is above the second floor.")
for item in [
    "The void is not a cold attic and shall not be freely ventilated with outdoor air.",
    "Maintain positive temperature and controlled humidity within the thermal and airtight boundary.",
    "Ventilation, electrical, ELV, PV and sensor routes may pass through the void if labelled and fully accessible.",
    "Roof penetrations shall use designed, sleeved and sealed details. Access hatches shall permit inspection, repair and equipment replacement without dismantling major assemblies.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-07 — Roof", "Provide a warm low-slope mono-pitch roof with an approximately 4° design slope, subject to final drainage and membrane-system requirements.")
for item in [
    "The roof is an independent watertight building element; the PV superstructure is not the primary roof or sole waterproofing.",
    "Coordinate structural deck, vapour/airtight layer, continuous insulation, falls, waterproofing, protection, primary drainage and emergency overflow.",
    "Keep outlets, penetrations, parapets and junctions accessible for inspection.",
    "PV supports shall connect only to calculated support zones or cast-in/installed anchors. Any connection across insulation requires a thermal break and fully warranted waterproof detail.",
    "Complete and test roof watertightness before PV superstructure installation.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-08 — PV Superstructure", "Provide a separate, maintainable structure over the warm roof to support solar panels, form the architectural roof edge and create a ventilated buffer.")
for item in [
    "Allow safe inspection, cleaning, isolation, removal and replacement without walking on panels or damaging roof waterproofing.",
    "Provide service walkways, ventilated under-panel clearance, controlled drainage, protected cable routes, earthing, isolators and required fire separation.",
    "Screen visible panel edges, framing and cables with a maintainable perimeter detail that does not obstruct ventilation or drainage.",
    "Coordinate a snow-removal/snow-melting provision where justified by climate and energy analysis.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-09 — External Stair Tower (Two-Storey Option)", "If the optional two-storey concept is proposed, provide an independent external stair tower for upper-floor access and, where required, technical access to the PV level.")
for item in [
    "Use its own foundation/support system and structural frame for self-weight, occupancy, wind and snow loads.",
    "Do not transfer vertical loads to the house façade, SIP bands, floor or roof. Use movement joints and weather-protected junctions.",
    "Provide code-compliant stairs, landings, guards, handrails, slip resistance and drainage.",
    "Any stabilising ties to the house require structural calculation, weather protection and thermal breaks.",
    "Coordinate lighting, access control, CCTV and optional landing snow-melting with the MEP design.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-10 — External Doors and Windows", "Treat glazing, frames, doors, insulated support profiles, thresholds, installation joints, drainage and maintenance access as one coordinated opening system.")
for item in [
    "Use high-performance double- or triple-chamber insulating glass units as required by calculation.",
    "Frames shall include thermal breaks, warm-edge spacers, drainage and serviceable hardware.",
    "Provide three-layer installation joints: internal airtight layer, middle insulation layer, external weather- and water-management layer.",
    "Use insulated sill/support profiles and warm thresholds. Foam alone is not an acceptable opening seal.",
    "Do not support frames directly on cold metal, concrete, the perimeter apron or unprepared framing.",
    "Thermally isolate fasteners, setting blocks and anchors. Glass units, seals, hardware and drains shall be replaceable.",
    "Select Ug, Uw/Ucw, solar factor, visible-light transmission, Low-E and solar-control coatings by orientation and energy model.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-11 — Thermal and Airtight Envelope", "Provide a continuous, inspectable boundary around every conditioned and protected volume.")
for item in [
    "The internal SIP enclosure is the principal thermal and airtight layer. Crawlspace and roof service void remain warm, accessible volumes.",
    "A glazed/SIP outer greenhouse buffer, if included, is thermally significant—not a cold decorative shell—and shall be analysed for temperature, humidity, overheating, condensation and purge ventilation.",
    "No breaks in insulation, uncontrolled air paths, cold thresholds or hidden condensation risks are permitted.",
    "All brackets, anchors, subframes and service penetrations crossing insulation require calculated thermal separation and sealed sleeves.",
    "Test and inspect the airtight layer before it is concealed by finishes. Supplier shall propose target air-tightness and test method for owner approval.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-12 — Serviceable Dry Hydronic Heated Floor", "Use a dry, lightweight and locally repairable system; avoid wet concrete screeds unless technically necessary and approved.")
for item in [
    "Place heating pipes in 40 mm routed XPS with aluminium heat-spreader plates.",
    "Cover with 18 mm MgO or two 10 mm layers, using elastomeric sealed joints, then the selected finish.",
    "Zone circuits by room and heat load; keep manifold cabinets accessible.",
    "Sleeve and protect pipes at penetrations and rubbing points.",
    "Enable local access by cutting sealant and lifting MgO boards without major floor demolition.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-13 — Serviceable Dry Hydronic Radiant Cooling Ceiling", "Use a dry hydronic ceiling system with mandatory humidity and dew-point control.")
for item in [
    "Place pipes in 40 mm routed XPS with heat-spreader plates; fix 10 mm MgO boards below on a calculated support system and seal joints elastically.",
    "Keep manifolds and joints accessible and coordinate with ventilation, lighting, fire requirements and service hatches.",
    "Do not operate below dew point. Provide automatic supply-temperature limitation, humidity sensors, temperature sensors and condensation protection.",
    "Design fixings for the full dead load of panels, pipes, water and finishes.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-14 — Internal Partitions, Doors and Windows", "Provide non-load-bearing, demountable modular partitions that support flexible planning, acoustic separation and local moisture resistance.")
for item in [
    "Preliminary concept: demountable 20–25 mm framing with gypsum board or MgO linings and mineral-wool acoustic infill; supplier shall verify practical thickness, stiffness, acoustic performance and service integration.",
    "Reinforce framing at televisions, cabinets and other concentrated loads.",
    "Provide dry-area and wet-area variants. Keep services accessible through removable panels or inspection openings.",
    "Door modules require reinforced frames, seals and compatible thresholds.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-15 — Internal Lining to External Walls", "Provide a serviceable internal lining/subframe to the SIP envelope and any greenhouse-buffer zones without compromising airtightness.")
for item in [
    "Maintain continuity of the internal airtight layer and seal electrical, ELV, sensor and fixing penetrations.",
    "Coordinate linings with window and door details and use moisture-resistant materials in wet/humid zones.",
    "Provide removable panels or inspection openings at manifolds, sensors, cable routes and ventilation components.",
]: add_bullet(doc, item)

add_section_intro(doc, "BDS-39 — Stormwater Drainage", "Collect and discharge water from the roof, PV superstructure, canopy, perimeter apron, paths, gates and hard landscaping without damaging foundations or creating ice hazards.")
for item in [
    "Coordinate stormwater with site drainage; do not combine with sanitary drainage without a specifically engineered and permitted solution.",
    "Use designed outlets, gutters, channels and discharge points. Grade aprons and paths toward collection or approved infiltration/discharge zones.",
    "Lay pipes at calculated falls with accessible cleanouts, silt traps and inspection chambers.",
    "Account for snow-melt water. Avoid conflicts with exterior cables, irrigation, gates and paths.",
    "Determine final discharge location from actual site conditions and local authority requirements.",
]: add_bullet(doc, item)

add_heading(doc, "6. Engineering, Documentation and Quality Requirements", 1)
add_heading(doc, "Design verification", 2)
for item in [
    "All structural members, panels, fasteners, joints, lifting points and temporary conditions shall be calculated for manufacture, transport, lifting, assembly and permanent use.",
    "Confirm design loads, climate data, seismic requirements, snow/wind loads, fire strategy, corrosion category, acoustic targets and destination-country codes before detailed design.",
    "Provide hygrothermal/condensation checks for critical envelope build-ups and thermal-bridge calculations for repeating and point connections.",
    "Clearly distinguish factory-installed, site-installed and locally sourced work.",
]: add_bullet(doc, item)
add_heading(doc, "Production and assembly information", 2)
for item in [
    "Issue coordinated IFC/shop drawings, module identification, connection schedules, tolerances, sealing sequences and torque requirements.",
    "Provide packing lists, weather protection, container/loading drawings, centre-of-gravity and lifting instructions.",
    "Provide a step-by-step assembly manual with labour/equipment requirements and QC hold points, including airtightness and waterproofing checks before concealment.",
    "Mark every service route, isolation point, cleanout, sensor, access panel and replaceable component.",
]: add_bullet(doc, item)
add_heading(doc, "Substitutions and deviations", 2)
add_text(doc, "The supplier may propose alternatives that improve manufacturability, cost or performance, but each alternative shall include technical data, effect on price and programme, effect on energy/fire/acoustic/durability performance, and a marked-up comparison with the stated requirement. No silent substitutions are permitted.")

add_heading(doc, "7. Information Required Before Detailed Design", 1)
for item in [
    "Site address, survey coordinates, levels, boundaries, utility locations and access constraints.",
    "Topographic/geodetic survey and geotechnical investigation, including groundwater and frost-depth data.",
    "Destination-country codes, permitting route, design climate and required certifications.",
    "Approved room schedule, target floor area, façade concept and glazing orientation.",
    "Confirmed utility capacities, electrical voltage/frequency and grid/PV connection requirements.",
    "Owner decisions on optional greenhouse buffer, hammam, car canopy, pellet boiler and snow-melting scope.",
    "Delivery route, container limits, crane access and maximum module dimensions/weights.",
]: add_bullet(doc, item)

add_callout(doc, "DOCUMENT STATUS", "This English brief communicates design intent for concept development and supplier quotation. It is not a construction drawing, structural calculation or permit document. Final design remains subject to coordinated engineering, actual site data and approvals by appropriately licensed local professionals.", fill="FFF8E8", accent=GOLD)

# Keep table rows and paragraphs visually stable.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True

doc.core_properties.title = "A-House 1 Core — Manufacturer Design Brief"
doc.core_properties.subject = "English supplier brief for prefabricated house manufacturers in China"
doc.core_properties.author = "A-House"
doc.core_properties.keywords = "prefabricated house, modular house, SIP, manufacturer brief, China"
doc.save(OUT)
print(OUT)
